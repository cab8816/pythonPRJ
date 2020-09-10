from django.contrib.auth.decorators import login_required
from django.contrib.auth.models import User
from django.db import models
from django.shortcuts import render
from docx import Document

from Myapp.models import PsyuanDetail, Biao4, Pszcy, Biao5, Biao72, Pingshenxxb, Xcpshcb71


def importpsymd(self, request, obj, change):  # 从word文件导入信息
    user1 = request.session.get('user1')
    print('session user1=' + str(user1))

    user1 = request.COOKIES.get('user1')
    print('cokie user1=' + str(user1))

    document = Document(obj.file)

    if obj.importtype == '0':  # (0, "评审通知文件"),
        pstzh = document.paragraphs[2].text
        if Pingshenxxb.objects.filter(pstzh=pstzh).count() < 1:  # 判断Pingshenxxb表中是否有这个评审通知的信息
            user = User.objects.filter(username=user1).first()  # 没有 就读入评审通知号 和 评审机构名称
            biao = Pingshenxxb(
                user=user,
                pstzh=document.paragraphs[2].text,
                jcjgmc=document.paragraphs[4].text,
            )
            biao.save()

            psxxb = obj.psxxb  # 拿到评审信息表的一行对象
            t = document.tables[0]
            if len(t.columns) == 5:  # 评审组成员表格
                print('生成  评审组成员表')
                for row in t.rows:
                    rowcells = row.cells
                    psname = rowcells[1].text
                    if psname != "姓 名":  # 去除表格标题的影响
                        psydtl = PsyuanDetail.objects.filter(name=psname).first()
                        biao = Pszcy(
                            psydtl=psydtl,
                            psyzc=rowcells[0].text,
                            psname=psname,

                            ziwuzicheng=rowcells[2].text,
                            gzdw=rowcells[3].text,
                            lxfs=rowcells[4].text,
                        )
                        biao.save()
                        biao.psxxbs.add(psxxb)  # 添加多到多的关系
    elif obj.importtype == '1':  # 读取评审系统导出的现场评审表
        for t in document.tables:
            table = t
            print('table len=' + str(len(t.columns)))
            psxxb = obj.psxxb
            if len(t.columns) == 11:  # 评审报告  表4 建议批准的检验检测能力表
                print('生成  表4')
                for row in table.rows:
                    rowcells = row.cells
                    lyname = rowcells[1].text
                    # if lyname != "领域":
                    biao = Biao4(
                        psxxb=psxxb,
                        lyxh=rowcells[0].text,
                        lyname=rowcells[1].text,
                        lbxh=rowcells[2].text,
                        lb=rowcells[3].text,
                        dxxh=rowcells[4].text,
                        duixiang=rowcells[5].text,
                        xmxh=rowcells[6].text,
                        csmc=rowcells[7].text,
                        yjbz=rowcells[8].text,
                        xzfw=rowcells[9].text,
                        sm=rowcells[10].text,
                    )
                    biao.save()


            elif len(t.columns) == 6:  # 建议批准的授权签字人
                print('生成  表5')
                for row in table.rows:
                    rowcells = row.cells

                    biao = Biao5(
                        psxxb=psxxb,
                        xuhao=rowcells[0].text,
                        name=rowcells[1].text,
                        ziwuzicheng=rowcells[3].text,
                        sqqzly=rowcells[4].text,
                        beizu=rowcells[5].text,
                    )
                    biao.save()

            elif len(t.columns) == 16:  # 表7.2 现场评审能力确认方式及确认结果一览表
                print('生成  表7.2')
                for row in table.rows:
                    rowcells = row.cells
                    xuhao = rowcells[0].text
                    if xuhao != "序号":  # 消除表头影响
                        biao = Biao72(
                            psxxb=psxxb,
                            xuhao=rowcells[0].text,
                            xmmc=rowcells[1].text,
                            yjbz=rowcells[2].text,
                            xmxh=rowcells[3].text,
                            csmc=rowcells[4].text,
                            # yjbztk = rowcells[0].text,
                            # xcsy = rowcells[0].text,
                            # nlyz = rowcells[0].text,
                            # clsh = rowcells[0].text,
                            # bdjg = rowcells[0].text,
                            # xcys = rowcells[0].text,
                            # xctw = rowcells[0].text,
                            # cyjlbg = rowcells[0].text,
                            # hcyq = rowcells[0].text,
                            # sfqr = rowcells[0].text,
                            # beizu = rowcells[0].text,
                        )

                        biao.save()




    elif obj.importtype == '2':  # (2, "资质认定评审员信息名单"),
        for t in document.tables:
            table = t
            if len(t.columns) == 5:  # 评审员名单
                for row in table.rows:
                    rowcells = row.cells
                    biao = PsyuanDetail(
                        name=rowcells[1].text,
                        gender=rowcells[2].text,
                        danwei=rowcells[3].text,
                        psybh=rowcells[4].text,
                    )
                    biao.save(force_insert=True)

    elif obj.importtype == '3':  # (2, "检验检测机构资质认定现场评审核查表7.1"),
        for t in document.tables:
            table = t
            psxxb = obj.psxxb  # 拿到评审信息表的一行对象
            if len(t.columns) == 4:  # 核查表
                for row in table.rows:
                    rowcells = row.cells
                    biao = Xcpshcb71(
                        psxxb=psxxb,
                        zhangbh=rowcells[0].text,
                        zhangmc=rowcells[1].text,
                        tkhao=rowcells[2].text,
                        psneirong=rowcells[3].text,
                        psjg='1',
                        pssm='',
                    )
                    biao.save()

    context = {'msg': '读取word文档成功'}
    return render(request, 'test.html', context)
